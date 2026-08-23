import io
import pdfplumber
import fitz  # PyMuPDF
import docx

from src.schemas.document import ParsedDocument

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10MB
MAX_PAGES = 20

def parse_document(file_content: bytes, filename: str) -> ParsedDocument:
    """
    Parses a document (PDF, DOCX, TXT) and returns the extracted text and metadata.
    """
    if len(file_content) > MAX_FILE_SIZE_BYTES:
        raise ValueError("File exceeds maximum allowed size of 10MB.")

    ext = filename.lower().split('.')[-1]

    if ext == "pdf":
        return _parse_pdf(file_content)
    elif ext == "docx":
        return _parse_docx(file_content)
    elif ext in ["txt", "md"]:
        return _parse_txt(file_content)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _parse_pdf(file_content: bytes) -> ParsedDocument:
    # Try pdfplumber first with layout preservation
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            page_count = len(pdf.pages)
            if page_count > MAX_PAGES:
                raise ValueError(f"PDF exceeds maximum allowed page count of {MAX_PAGES}.")
            
            text_parts = []
            for page in pdf.pages:
                extracted = page.extract_text(layout=True)
                if extracted:
                    text_parts.append(extracted)
            
            raw_text = "\n".join(text_parts)
            
            # If text is extracted successfully and doesn't look garbled, return it
            if not _is_text_garbled(raw_text):
                return ParsedDocument(
                    raw_text=raw_text.strip(),
                    source_format="pdf",
                    page_count=page_count,
                    extraction_method="pdfplumber"
                )
    except ValueError as ve:
        raise ve
    except Exception:
        # Ignore other pdfplumber exceptions and fall back to PyMuPDF
        pass

    # Fallback to PyMuPDF (fitz)
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        if doc.needs_pass:
            raise ValueError("PDF is password-protected.")
            
        page_count = len(doc)
        if page_count > MAX_PAGES:
            raise ValueError(f"PDF exceeds maximum allowed page count of {MAX_PAGES}.")
            
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
            
        raw_text = "\n".join(text_parts)
        if not raw_text.strip():
            raise ValueError("Failed to extract any text from the PDF.")
            
        return ParsedDocument(
            raw_text=raw_text.strip(),
            source_format="pdf",
            page_count=page_count,
            extraction_method="pymupdf"
        )
    except ValueError as ve:
        raise ve
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def _parse_docx(file_content: bytes) -> ParsedDocument:
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text_parts = [para.text for para in doc.paragraphs]
        raw_text = "\n".join(text_parts)
        
        return ParsedDocument(
            raw_text=raw_text.strip(),
            source_format="docx",
            page_count=1,  # python-docx doesn't easily support page counting
            extraction_method="python-docx"
        )
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def _parse_txt(file_content: bytes) -> ParsedDocument:
    try:
        raw_text = file_content.decode("utf-8")
        return ParsedDocument(
            raw_text=raw_text.strip(),
            source_format="txt",
            page_count=1,
            extraction_method="txt_decode"
        )
    except UnicodeDecodeError:
        try:
            # Fallback for older encodings
            raw_text = file_content.decode("latin-1")
            return ParsedDocument(
                raw_text=raw_text.strip(),
                source_format="txt",
                page_count=1,
                extraction_method="txt_decode"
            )
        except Exception as e:
            raise ValueError(f"Failed to parse TXT: {str(e)}")


def _is_text_garbled(text: str) -> bool:
    """
    Heuristic to determine if the extracted text is garbled or malformed.
    """
    if not text.strip():
        return True
    
    # Heuristic 1: Low ratio of alphanumeric characters (excluding whitespace)
    # layout=True uses many spaces, so we must ignore them for the ratio
    text_no_space = "".join(text.split())
    if not text_no_space:
        return True
        
    alnum_count = sum(c.isalnum() for c in text_no_space)
    if alnum_count / len(text_no_space) < 0.5:
        return True
        
    # Heuristic 2: Excessive single-character words (often happens with bad font maps)
    words = text.split()
    if not words:
        return True
        
    single_char_words = sum(1 for w in words if len(w) == 1 and w.isalpha())
    if single_char_words / len(words) > 0.5:
        return True
        
    return False
