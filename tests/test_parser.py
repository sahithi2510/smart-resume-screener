import pytest
import io
from fpdf import FPDF
import docx
from src.services.parser import parse_document, MAX_FILE_SIZE_BYTES, MAX_PAGES

def generate_single_column_pdf() -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, text="Work Experience", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(200, 10, text="Software Engineer at Tech Corp", new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())

def generate_multi_column_pdf() -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Left column
    pdf.set_xy(10, 10)
    pdf.cell(80, 10, text="Work Experience")
    pdf.set_xy(10, 20)
    pdf.cell(80, 10, text="Software Engineer")

    # Right column
    pdf.set_xy(100, 10)
    pdf.cell(80, 10, text="Skills")
    pdf.set_xy(100, 20)
    pdf.cell(80, 10, text="Python, FastAPI")
    
    return bytes(pdf.output())

def generate_docx() -> bytes:
    doc = docx.Document()
    doc.add_paragraph("This is a sample DOCX resume.")
    doc.add_paragraph("Skills: Python, SQL.")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def test_parse_single_column_pdf():
    content = generate_single_column_pdf()
    parsed = parse_document(content, "resume.pdf")
    
    assert parsed.source_format == "pdf"
    assert "Work Experience" in parsed.raw_text
    assert "Tech Corp" in parsed.raw_text
    assert parsed.extraction_method == "pdfplumber"
    assert parsed.page_count == 1

def test_parse_multi_column_pdf():
    content = generate_multi_column_pdf()
    parsed = parse_document(content, "resume.pdf")
    
    assert parsed.source_format == "pdf"
    assert "Work Experience" in parsed.raw_text
    assert "Skills" in parsed.raw_text
    assert "Python, FastAPI" in parsed.raw_text
    assert parsed.extraction_method == "pdfplumber"
    
def test_parse_docx():
    content = generate_docx()
    parsed = parse_document(content, "resume.docx")
    
    assert parsed.source_format == "docx"
    assert "sample DOCX resume" in parsed.raw_text
    assert parsed.extraction_method == "python-docx"

def test_parse_txt():
    content = b"Simple text resume\nPython developer"
    parsed = parse_document(content, "resume.txt")
    
    assert parsed.source_format == "txt"
    assert "Simple text resume" in parsed.raw_text
    assert parsed.extraction_method == "txt_decode"

def test_file_size_limit():
    content = b"0" * (MAX_FILE_SIZE_BYTES + 1)
    with pytest.raises(ValueError, match="exceeds maximum allowed size"):
        parse_document(content, "large.txt")

def test_page_count_limit():
    pdf = FPDF()
    for _ in range(MAX_PAGES + 1):
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, text="Page")
    content = bytes(pdf.output())
    
    with pytest.raises(ValueError, match="exceeds maximum allowed page count"):
        parse_document(content, "toolong.pdf")
        
def test_corrupt_pdf():
    content = b"%PDF-1.4\n%This is a corrupt pdf that will fail parsing\n0 0 0 0"
    with pytest.raises(ValueError, match="Failed to parse PDF|Failed to extract"):
        parse_document(content, "corrupt.pdf")
